import os
import csv
from datetime import datetime
import PyPDF2
import docx
import pypdf
import pickle
from utils.config_loader import load_config


def read_txt(file_path):
    """Reads a plain text file, returns string content."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_pdf(file_path):
    """Extracts text from a PDF file using PyPDF2, returns string content."""
    text = []
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text(
                extraction_mode="layout", layout_mode_scale_weight=0.1
            )
            if page_text:
                text.append(page_text)
    return "\n".join(text)


def read_docx_file(file_path):
    """Extracts text from a .docx file using python-docx, returns string content."""
    doc = docx.Document(file_path)
    full_text = []
    # Extract from top-level paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)

    return "\n".join(full_text)


def extract_text(file_path):
    """
    Determines the file extension and uses the appropriate function
    to extract text. Returns the extracted text or an empty string on unsupported types.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        return read_txt(file_path)
    elif ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx_file(file_path)
    else:
        # You can log this if needed. For now, just return empty to skip.
        return ""


def generate_csv(base_folder, output_csv):
    """
    base_folder structure (example):
        base_folder/
            Resumes/
            CoverLetters/
            Essays/
    Each subfolder is considered 'doc_type'.
    We detect file extension (pdf, docx, txt) and extract text accordingly.
    """
    with open(output_csv, mode="w", newline="", encoding="utf-8") as csv_file:
        writer = csv.writer(csv_file)
        # CSV header
        writer.writerow(["id", "content", "doc_type", "created_date"])

        # Traverse subfolders
        for doc_type in os.listdir(base_folder):
            doc_type_path = os.path.join(base_folder, doc_type)

            if os.path.isdir(doc_type_path):
                for filename in os.listdir(doc_type_path):
                    file_path = os.path.join(doc_type_path, filename)

                    # Extract text based on file type
                    content = extract_text(file_path)

                    if content.strip():  # If there's any text
                        doc_id = filename.replace(" ", "_")
                        created_date = datetime.now().strftime("%Y-%m-%d")

                        writer.writerow(
                            [
                                doc_id,  # "doc001", or "My_Resume.pdf", etc.
                                content,  # extracted text
                                doc_type,  # "Resumes", "CoverLetters", etc.
                                created_date,
                            ]
                        )


# File: resume_parser_project/config/label_config.py

import json


def load_flat_labels(json_path: str) -> dict[str, dict]:
    """
    Loads a flattened JSON containing both top-level section headings and sub-labels.
    Returns a dictionary mapping every lowercased synonym -> metadata dict.

    Example of returned dictionary structure:
      {
        "professional experience": {
          "label_type": "SECTION_HEADING",
          "canonical_label": "PROFESSIONAL EXPERIENCE",
          "belongs_to": None
        },
        "work experience": {
          "label_type": "SECTION_HEADING",
          "canonical_label": "PROFESSIONAL EXPERIENCE",
          "belongs_to": None
        },
        "career history": {
          "label_type": "SECTION_HEADING",
          "canonical_label": "PROFESSIONAL EXPERIENCE",
          "belongs_to": None
        },
        "company line": {
          "label_type": "SUB_LABEL",
          "canonical_label": "Company Line",
          "belongs_to": "PROFESSIONAL EXPERIENCE"
        },
        ...
      }

    :param json_path: Path to your flattened JSON file (e.g., headings_and_sublabels.json).
    :return: A dict where keys are synonyms (lowercase), and values are dictionaries with metadata.
    """
    # with open(json_path, "r", encoding="utf-8") as f:
    #     data = json.load(f)
    data = load_config(json_path)
    flat_map: dict[str, dict] = {}

    for entry in data["labels"]:
        label_type = entry["label_type"]
        canonical_label = entry["canonical_label"]
        synonyms = entry.get("synonyms", [])
        belongs_to = entry.get("belongs_to", None)  # For sub-labels only

        # Map the canonical label itself
        flat_map[canonical_label.lower()] = {
            "label_type": label_type,
            "canonical_label": canonical_label,
            "belongs_to": belongs_to,
        }

        # Map each synonym
        for syn in synonyms:
            flat_map[syn.lower()] = {
                "label_type": label_type,
                "canonical_label": canonical_label,
                "belongs_to": belongs_to,
            }

    return flat_map


def save_object_to_pickle(object: str, file_path: str) -> None:
    """Save transcription object to a pickle file."""
    with open(file_path, "wb") as f:
        pickle.dump(object, f)
    print("pickle file saved")


def load_object_from_pickle(file_path: str) -> any:
    """Load transcription object from a pickle file."""
    with open(file_path, "rb") as f:
        return pickle.load(f)


def build_synonyms_map(flat_map: dict[str, dict]) -> dict[str, list[str]]:
    """
    Groups all known surface forms (canonical label + synonyms)
    under their canonical label keys.
    """
    synonyms_map = {}
    for text_lower, label_info in flat_map.items():
        canonical_label = label_info["canonical_label"]
        if canonical_label not in synonyms_map:
            synonyms_map[canonical_label] = set()  # use a set to avoid duplicates
        synonyms_map[canonical_label].add(
            text_lower
        )  # text_lower is either the c-label or a synonym

    # Convert sets to lists
    return {k: list(v) for k, v in synonyms_map.items()}
